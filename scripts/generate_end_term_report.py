from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_PATH = Path("/Users/a1234/Documents/AITU/SRE/Assignment4-5/foodrush/FoodRush_End_Term_Report.docx")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 20, RGBColor(0x1F, 0x1F, 0x1F)),
        ("Heading 1", 14, RGBColor(0x1F, 0x4E, 0x78)),
        ("Heading 2", 11.5, RGBColor(0x1F, 0x4E, 0x78)),
        ("Heading 3", 10.5, RGBColor(0x2F, 0x5D, 0x8A)),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    quote = doc.styles["Intense Quote"]
    quote.font.name = "Arial"
    quote.font.size = Pt(10)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FoodRush End-Term Report | Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    add_page_number(p)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("END TERM PROJECT")
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Comprehensive SRE Implementation for a Distributed Microservices System"
    )
    run.font.name = "Arial"
    run.font.size = Pt(12)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("Student", "Sagyntai Aikyn"),
        ("Group", "SE-2429"),
        ("Project", "FoodRush"),
        ("Date", "16 May 2026"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], "D9EAF7")

    doc.add_paragraph(
        "Screenshot note: before final submission, mask public IP addresses, usernames, ports, browser URLs, project identifiers, and all secrets."
    )


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_placeholder(doc, text):
    doc.add_paragraph(f"[Insert Screenshot: {text}]", style="Intense Quote")


def add_service_block(doc, name, description, screenshot_text):
    add_heading(doc, name, level=3)
    doc.add_paragraph(description)
    add_placeholder(doc, screenshot_text)


def build_report():
    doc = Document()
    set_styles(doc)
    add_footer(doc)
    add_title(doc)

    add_heading(doc, "1. Title")
    doc.add_paragraph(
        "End-to-End Implementation of Site Reliability Engineering Practices in a Multi-Orchestrated Microservices Infrastructure Using Docker Swarm, Kubernetes, Terraform, and Ansible"
    )

    add_heading(doc, "2. Abstract")
    doc.add_paragraph(
        "This project presents a complete implementation of Site Reliability Engineering principles applied to FoodRush, a distributed microservices-based food delivery platform. The final system integrates containerization, multiple orchestration layers, infrastructure provisioning, deployment automation, monitoring, alerting, incident validation, and capacity-oriented analysis."
    )
    doc.add_paragraph(
        "The architecture contains six backend microservices, a frontend, PostgreSQL databases, Prometheus, and Grafana. Docker Compose was used as the baseline deployment model, while Docker Swarm and Kubernetes were implemented to demonstrate orchestration beyond single-node execution. Terraform was used to provision the cloud environment, and Ansible was used to automate host preparation and application deployment."
    )
    doc.add_paragraph(
        "Observability was implemented through Prometheus metrics collection, Grafana dashboards, and alert rules. A controlled incident affecting order-service database connectivity was simulated to validate degraded-mode behavior, monitoring visibility, and recovery. Capacity analysis was also performed to identify bottlenecks and define practical scaling recommendations."
    )

    add_heading(doc, "3. Objectives")
    doc.add_paragraph(
        "The main objective of the end-term project is to demonstrate SRE practices across the full system lifecycle, from deployment and provisioning to reliability validation and operational analysis."
    )
    add_numbered(
        doc,
        [
            "Design and deploy a distributed microservices architecture with at least six backend services.",
            "Implement multiple orchestration approaches using Docker Swarm and Kubernetes.",
            "Provision cloud infrastructure using Terraform.",
            "Automate configuration and deployment using Ansible.",
            "Define practical SLIs and SLOs based on collected service metrics.",
            "Implement monitoring and alerting with Prometheus and Grafana.",
            "Simulate a realistic service incident and validate recovery.",
            "Document operational findings and postmortem conclusions.",
            "Perform load-based capacity analysis and propose scaling strategies.",
        ],
    )

    add_heading(doc, "4. System Overview")
    doc.add_paragraph(
        "FoodRush is a scalable microservices platform designed as an SRE demonstration environment. The final version includes six backend services, one frontend, dedicated monitoring components, and PostgreSQL persistence. To make later screenshot insertion easier, every major component is described in its own block below."
    )

    add_service_block(
        doc,
        "4.1 auth-service",
        "auth-service handles authentication, token verification, and identity-related operations. It is responsible for user login, token issuance, and role-aware authentication checks used by the rest of the platform.",
        "auth-service code, routes, or auth-service API evidence",
    )
    add_service_block(
        doc,
        "4.2 product-service",
        "product-service provides restaurant, menu, and product catalog functionality. It is also the most important service for the capacity section because heavy read traffic exposed it as the main bottleneck under stress.",
        "product-service code, menu API, or product-service implementation evidence",
    )
    add_service_block(
        doc,
        "4.3 order-service",
        "order-service handles order placement and lifecycle logic. It also exposes degraded-mode health behavior, which made it the natural target for the incident simulation section.",
        "order-service code or order-service API/health evidence",
    )
    add_service_block(
        doc,
        "4.4 user-service",
        "user-service manages customer profiles, saved addresses, and user-related operations required by the frontend and order workflow.",
        "user-service code or user-service API evidence",
    )
    add_service_block(
        doc,
        "4.5 chat-service",
        "chat-service supports order-related communication and support chat rooms. It extends the platform beyond core CRUD behavior and demonstrates service-to-service integration.",
        "chat-service code or chat-service API evidence",
    )
    add_service_block(
        doc,
        "4.6 payment-service",
        "payment-service handles payment creation, lookup, and payment-status operations. It was added as the sixth backend service to satisfy the final project microservices scope.",
        "payment-service code or payment-service API evidence",
    )
    add_service_block(
        doc,
        "4.7 Frontend",
        "The frontend is served through Nginx and acts as the user-facing entry point to the platform. It integrates all backend services through reverse-proxied API routes.",
        "frontend UI or frontend source structure",
    )
    add_service_block(
        doc,
        "4.8 Supporting Components",
        "The system also includes PostgreSQL databases for service persistence, Prometheus for metrics scraping, and Grafana for visualization. These components are essential to the operational, not just functional, side of the platform.",
        "repository structure or supporting components overview",
    )

    add_heading(doc, "5. Deployment and Containerization")
    doc.add_paragraph(
        "Before orchestration and automation layers were added, the project was first organized as a containerized application. This section groups the core deployment artifacts in one place so that implementation screenshots can be inserted directly after each relevant block."
    )
    add_heading(doc, "5.1 Docker Compose", level=2)
    doc.add_paragraph(
        "Docker Compose is the baseline deployment model for FoodRush. It defines service-to-service networking, service startup order, restart behavior, health checks, database containers, monitoring services, and exposed frontend/API entry points."
    )
    add_placeholder(doc, "docker-compose.yml")
    add_placeholder(doc, "docker compose ps")

    add_heading(doc, "5.2 Service Dockerfiles", level=2)
    doc.add_paragraph(
        "Each backend service is packaged with its own Dockerfile. These files standardize runtime dependencies, application startup commands, and the container structure used later by Compose, Swarm, and Kubernetes."
    )
    add_placeholder(doc, "backend service Dockerfile")
    add_placeholder(doc, "payment-service Dockerfile or another representative Dockerfile")

    add_heading(doc, "5.3 Frontend Containerization", level=2)
    doc.add_paragraph(
        "The frontend is served through Nginx. Its container image packages static assets and applies a custom Nginx configuration so the web interface can proxy or expose the platform correctly."
    )
    add_placeholder(doc, "frontend/Dockerfile")
    add_placeholder(doc, "nginx.conf")
    add_placeholder(doc, "frontend running through Docker Compose")

    add_heading(doc, "6. Multi-Orchestration Architecture")
    add_heading(doc, "6.1 Docker Swarm", level=2)
    doc.add_paragraph(
        "Docker Swarm was implemented to demonstrate service replication, overlay networking, and fast orchestration on top of a simple manager-node setup. It provides a clear step forward from plain Compose by introducing service-level abstraction and replicated workloads."
    )
    add_bullets(
        doc,
        [
            "Simple clustering model for containerized services",
            "Service replication for stateless application components",
            "Fast deployment and removal through stack commands and helper scripts",
        ],
    )
    add_placeholder(doc, "docker service ls")
    add_placeholder(doc, "docker stack services foodrushswarm")
    add_placeholder(doc, "docker stack ps foodrushswarm or frontend through Swarm")

    add_heading(doc, "6.2 Kubernetes", level=2)
    doc.add_paragraph(
        "Kubernetes was implemented to demonstrate a more advanced orchestration model. Compared with Swarm, it provides stronger workload abstraction, namespace-level organization, and a more production-oriented control model."
    )
    add_bullets(
        doc,
        [
            "Declarative workload definitions through manifests",
            "Deployments, Services, and StatefulSets for structured orchestration",
            "Namespace-based isolation for the application and monitoring stack",
            "Operational readiness for future scaling policies",
        ],
    )
    add_placeholder(doc, "kubectl get pods -n foodrush")
    add_placeholder(doc, "kubectl get svc -n foodrush")
    add_placeholder(doc, "kubectl get deployments,statefulsets -n foodrush")

    add_heading(doc, "6.3 Justification", level=2)
    doc.add_paragraph(
        "Using both Docker Swarm and Kubernetes makes the infrastructure section stronger because it demonstrates comparative orchestration rather than only one deployment path. Swarm provides speed and simplicity, while Kubernetes demonstrates stronger abstraction and operational depth. Together they show flexibility in deployment strategy and a broader understanding of SRE infrastructure design."
    )
    add_placeholder(doc, "Architecture diagram or orchestration comparison evidence")

    add_heading(doc, "7. Infrastructure Provisioning with Terraform")
    doc.add_paragraph(
        "Terraform was used to provision the cloud infrastructure required for final deployment. The configuration creates the compute instance, required access rules, and the output values needed for deployment verification."
    )
    add_bullets(
        doc,
        [
            "Provision cloud infrastructure in Google Cloud",
            "Configure network access rules required for the platform",
            "Ensure reproducibility through declarative infrastructure code",
        ],
    )
    doc.add_paragraph("Key benefits of the Terraform layer:")
    add_bullets(
        doc,
        [
            "Declarative configuration instead of manual console setup",
            "Version-controlled infrastructure changes",
            "Repeatable provisioning workflow",
        ],
    )
    add_placeholder(doc, "terraform init")
    add_placeholder(doc, "terraform plan")
    add_placeholder(doc, "terraform apply")
    add_placeholder(doc, "terraform output with sensitive fields masked")

    add_heading(doc, "8. Configuration Management with Ansible")
    doc.add_paragraph(
        "Ansible automates host preparation and deployment on the target VM. The final playbook flow installs Docker, prepares directories, renders the environment file, validates configuration, transfers the project, removes conflicting legacy containers, and starts the application stack."
    )
    add_bullets(
        doc,
        [
            "Install Docker and required host dependencies",
            "Configure the runtime environment for the application",
            "Deploy the stack using Docker Compose or Docker Swarm",
            "Support monitoring-ready deployment in the same automated workflow",
        ],
    )
    doc.add_paragraph("Main task groups covered by Ansible:")
    add_numbered(
        doc,
        [
            "Install dependencies and prepare the host.",
            "Configure directories and environment files.",
            "Deploy the application stack.",
            "Validate resulting service state.",
        ],
    )
    add_placeholder(doc, "ansible all -m ping")
    add_placeholder(doc, "ansible-playbook playbooks/site.yml with successful PLAY RECAP")
    add_placeholder(doc, "successful docker compose ps after Ansible deployment")

    add_heading(doc, "9. Monitoring, Alerting, and SLI/SLO")
    doc.add_paragraph(
        "Prometheus and Grafana form the observability layer of the project. Prometheus scrapes metrics from all six backend services, while Grafana visualizes service health and operational trends needed for reliability analysis."
    )
    doc.add_paragraph("Collected metrics include:")
    add_numbered(
        doc,
        [
            "CPU usage",
            "Memory usage",
            "Request rate",
            "5xx error rate",
            "Service availability and healthy target count",
            "P95 latency",
        ],
    )
    doc.add_paragraph(
        "The final monitoring layer also supports concrete SLI and SLO definitions tied to implemented dashboards and Prometheus rules."
    )
    add_heading(doc, "9.1 SLI/SLO Definition", level=2)
    doc.add_paragraph("Availability SLI/SLO:")
    doc.add_paragraph(
        "Availability was defined as the percentage of healthy targets and successful service health exposure. The target threshold was set to at least 99%."
    )
    add_placeholder(doc, "availability graph or query evidence")
    doc.add_paragraph("Latency SLI/SLO:")
    doc.add_paragraph(
        "P95 latency was used as the main response-time indicator, with a target threshold of 200 ms or lower during normal traffic."
    )
    add_placeholder(doc, "P95 latency graph or query evidence")
    doc.add_paragraph("Error Rate SLI/SLO:")
    doc.add_paragraph(
        "The ratio of 5xx responses to total requests was used as the error-rate indicator, with a target threshold below 1%."
    )
    add_placeholder(doc, "5xx error rate graph or query evidence")

    add_heading(doc, "9.2 Monitoring Evidence", level=2)
    add_placeholder(doc, "Prometheus targets page")
    add_placeholder(doc, "Prometheus alerts page")
    add_placeholder(doc, "Grafana healthy dashboard")

    add_heading(doc, "10. Incident Simulation and Postmortem")
    doc.add_paragraph(
        "A controlled incident was simulated by intentionally breaking ORDER_DATABASE_URL for order-service. This scenario was selected because it tests whether the service degrades in a visible and diagnosable way instead of failing silently."
    )
    add_heading(doc, "10.1 Incident Overview", level=2)
    doc.add_paragraph(
        "The simulated incident affected order-service database connectivity. order-service remained reachable, but order-related functionality degraded because database access was lost. The event was visible through application health output, logs, and monitoring metrics."
    )
    add_placeholder(doc, "healthy baseline before incident")
    add_heading(doc, "10.2 Customer Impact", level=2)
    doc.add_paragraph(
        "Users could still reach the application, but order-related operations experienced degraded behavior. This created a realistic partial outage rather than a total platform shutdown, which is useful for SRE analysis because it exercises degraded-mode observability."
    )
    add_placeholder(doc, "frontend during incident")
    add_heading(doc, "10.3 Root Cause Analysis", level=2)
    doc.add_paragraph(
        "The root cause was deliberate misconfiguration of ORDER_DATABASE_URL. Because order-service depends on successful database access for normal order processing, the service entered degraded mode and exposed the failure through its health endpoint and metrics."
    )
    add_placeholder(doc, "broken ORDER_DATABASE_URL or incident configuration evidence")
    add_heading(doc, "10.4 Detection and Response", level=2)
    doc.add_paragraph(
        "The incident was detected through multiple sources: the /health endpoint returned degraded state, order-service logs showed database-related failures, Prometheus captured changed metrics, and Grafana visualized the degraded period and subsequent recovery."
    )
    add_placeholder(doc, "degraded /health response")
    add_placeholder(doc, "order-service logs during incident")
    add_placeholder(doc, "Prometheus alert during incident")
    add_placeholder(doc, "Grafana during incident")
    add_heading(doc, "10.5 Resolution Summary", level=2)
    doc.add_paragraph(
        "Recovery was performed by restoring the correct ORDER_DATABASE_URL value and restarting order-service. After recovery, the health endpoint returned to normal and monitoring graphs returned to baseline."
    )
    add_placeholder(doc, "recovered /health response")
    add_placeholder(doc, "Grafana after recovery")
    add_heading(doc, "10.6 Lessons Learned", level=2)
    add_bullets(
        doc,
        [
            "Configuration errors remain a practical source of service degradation.",
            "Degraded-mode exposure is valuable because it preserves visibility instead of hiding failures.",
            "Monitoring and logs together provide better diagnosis than either source alone.",
        ],
    )

    add_heading(doc, "11. Automation and Capacity Planning")
    doc.add_paragraph(
        "Automation in the final project is not limited to deployment. It also includes health checks, restart behavior, environment validation, and faster troubleshooting. Capacity planning extends this by showing how the system behaves under increased load and where the first real bottleneck appears."
    )
    add_bullets(
        doc,
        [
            "Automated deployment through Docker Compose, Terraform, and Ansible",
            "Health checks and restart policies for improved operational resilience",
            "Environment validation and operational log inspection scripts",
            "Load-based observation of performance under higher concurrency",
        ],
    )
    doc.add_paragraph(
        "The main capacity finding was that product-service became the first significant bottleneck under heavier read traffic. The observed technical cause was connection pool pressure in the database access layer."
    )
    add_placeholder(doc, "environment validation script output")
    add_placeholder(doc, "log inspection or automation script evidence")
    add_placeholder(doc, "load baseline dashboard")
    add_placeholder(doc, "heavy load dashboard")
    add_placeholder(doc, "product-service pool timeout log")

    add_heading(doc, "12. Results and Operational Assessment")
    doc.add_paragraph(
        "The final system demonstrates practical multi-orchestrated deployment, infrastructure automation, monitoring, alerting, and recoverable incident behavior. From an operational perspective, the project is not only deployable but also observable and diagnosable."
    )
    add_heading(doc, "12.1 Microservices Architecture", level=2)
    doc.add_paragraph(
        "The final backend consists of six services, and the platform is no longer limited to a minimal proof of concept."
    )
    add_placeholder(doc, "evidence for six-service architecture")
    add_heading(doc, "12.2 Deployment Models", level=2)
    doc.add_paragraph(
        "Compose, Swarm, and Kubernetes were all implemented and validated, which significantly strengthens the operational maturity of the project."
    )
    add_placeholder(doc, "deployment-model comparison evidence")
    add_heading(doc, "12.3 Monitoring and Incident Readiness", level=2)
    doc.add_paragraph(
        "Monitoring and alerting were not only configured but also exercised through a real incident simulation and post-recovery validation."
    )
    add_placeholder(doc, "monitoring and incident-readiness evidence")
    add_heading(doc, "12.4 Capacity Readiness", level=2)
    doc.add_paragraph(
        "Heavy-load testing exposed a concrete system bottleneck, which is stronger evidence than a generic claim that the platform is scalable."
    )
    add_placeholder(doc, "capacity analysis evidence")

    add_heading(doc, "13. Conclusion")
    doc.add_paragraph(
        "The end-term project demonstrates a full SRE-oriented implementation around the FoodRush platform. The final system includes distributed services, multiple orchestration models, infrastructure as code, deployment automation, monitoring, alerting, incident validation, and capacity analysis."
    )
    doc.add_paragraph(
        "The most important engineering outcome is that the project now behaves like an operable system rather than only an academic application. It can be provisioned reproducibly, deployed automatically, monitored continuously, and analyzed through real incidents and performance behavior. At the same time, the capacity analysis revealed a concrete bottleneck in product-service database access, which gives a practical basis for future scaling and optimization."
    )
    add_placeholder(doc, "final healthy system state after all testing")

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    print(build_report())
